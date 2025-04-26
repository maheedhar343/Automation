import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import load_workbook
import os
import subprocess
import sys
import re
from styling import lighten_color, set_cell_shading, set_cell_margins, set_cell_border, set_table_borders, format_text_with_bullets

# Get file paths from command-line arguments
if len(sys.argv) != 3:
    print("Usage: python generate_document.py <excel_path> <template_path>")
    sys.exit(1)

excel_path = sys.argv[1]  # Single Excel sheet
template_path = sys.argv[2]
output_file = "Final_output.docx"  # Output file path

# Step 1: Load Excel and extract severity colors
try:
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
except FileNotFoundError:
    print(f"Error: Excel file '{excel_path}' not found.")
    sys.exit(1)

# Get headers and map severity column
headers = [cell.value for cell in ws[1] if cell.value is not None]
header_lookup = {str(val).strip().lower(): idx for idx, val in enumerate(headers)}
if "severity" not in header_lookup:
    print("Error: 'Severity' column not found in the Excel header.")
    sys.exit(1)
severity_col_index = header_lookup["severity"] + 1  # 1-based for openpyxl

# Map row index to severity fill color (in hex)
severity_colors = {}
for row in ws.iter_rows(min_row=2):
    row_idx = row[0].row
    severity_cell = ws.cell(row=row_idx, column=severity_col_index)
    fill = severity_cell.fill
    if fill and fill.fgColor and fill.fgColor.type == 'rgb':
        severity_colors[row_idx - 2] = fill.fgColor.rgb[-6:]  # Get RRGGBB from AARRGGBB

# Step 2: Load Excel data with pandas
try:
    df = pd.read_excel(excel_path)
except FileNotFoundError:
    print(f"Error: Excel file '{excel_path}' not found.")
    sys.exit(1)

# Normalize Excel column names for case-insensitive matching
excel_columns = df.columns.tolist()
excel_columns_normalized = {str(col).strip().lower(): col for col in excel_columns}

# Identify columns from "Proof of Concept" onwards
poc_index = None
for idx, col in enumerate(excel_columns):
    if str(col).strip().lower() == "proof of concept":
        poc_index = idx
        break

if poc_index is None:
    print("Error: 'Proof of Concept' column not found in the Excel sheet.")
    sys.exit(1)

# Get all columns from "Proof of Concept" to the end
additional_columns = excel_columns[poc_index:]

# Step 3: Load Word template
try:
    doc = Document(template_path)
except FileNotFoundError:
    print(f"Error: Word template '{template_path}' not found.")
    sys.exit(1)

# Set page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Step 4: Find matching table in template
tables = doc.tables
matching_table = None
for table in tables:
    if table.cell(0, 0).text.strip().lower() in excel_columns_normalized:
        matching_table = table
        break

if not matching_table:
    print("Error: No matching table found in the template.")
    sys.exit(1)

# Save trailing content
tbl_elm = matching_table._element
parent = tbl_elm.getparent()
index_in_body = list(parent).index(tbl_elm)
following_elements = list(parent)[index_in_body + 1:]
parent.remove(tbl_elm)

# Get headers from matching table
row_headers = [row.cells[0].text.strip() for row in matching_table.rows]

# Step 5: Generate content per row
for idx, row in df.iterrows():
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_para.add_run(f"Table {idx + 1}")
    run.font.size = Pt(16)
    run.bold = True
    
    # Get severity color for this row
    severity_hex = severity_colors.get(idx)
    light_severity_hex = lighten_color(severity_hex) if severity_hex else None
    
    # Count rows needed (excluding POC, but we'll add it as the last row)
    table_rows = sum(1 for header in row_headers if header.lower() != "proof of concept")
    table_rows += 1  # Add one row for POC and subsequent columns
    
    # Create a table
    table = doc.add_table(rows=table_rows, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.width = Inches(6.5)  # Fit page width (8.5 - 1 - 1 = 6.5 inches)
    
    # Fill the table with data (excluding POC)
    row_index = 0
    
    for i, header in enumerate(row_headers):
        if header.lower() == "proof of concept":
            continue  # Skip POC since we'll handle it in the last row
            
        # Add data to the table cell
        cell = table.cell(row_index, 0)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify alignment for all cells
        paragraph.paragraph_format.line_spacing = 1.0
        
        # Add padding
        set_cell_margins(cell, margin_value=100)
        
        # Apply background color to the first three rows
        if i < 3 and severity_hex:
            if i < 2:  # First two rows: original severity color
                set_cell_shading(cell, severity_hex)
            elif i == 2:  # Third row: lighter color
                set_cell_shading(cell, light_severity_hex)
        
        # Normalize header for matching
        header_normalized = header.strip().lower()
        excel_header = excel_columns_normalized.get(header_normalized, header)
        
        # Add the content (modify first two rows to show only data)
        text_value = str(row.get(excel_header, '')).strip()
        print(f"Row {idx + 1}, Table Row {i + 1}, Header '{header}', Excel Header '{excel_header}': Text Value = '{text_value}'")  # Debug print
        if i < 2:  # First two rows: show only data, no header
            run = paragraph.add_run(text_value)
            run.font.size = Pt(11) if i == 0 else Pt(14)  # First row: 11 pt, second row: 14 pt
            run.bold = True
        else:  # Other rows: show header and value
            # Add column name with text color
            run = paragraph.add_run(f"{header}:")
            run.font.size = Pt(11)
            run.bold = True
            if severity_hex:
                # Convert hex color to RGB and apply as text color
                r = int(severity_hex[0:2], 16)
                g = int(severity_hex[2:4], 16)
                b = int(severity_hex[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            
            # Add a line break before the data only for rows 9 and beyond (indices 8+)
            if i >= 8:  # 9th row and beyond (indices 8, 9, etc.)
                cell.add_paragraph()  # Add a new paragraph for visual separation
            # Note: For rows 0 to 7 (1st to 8th rows), no line break is added before the data
            
            # Apply bullet points for 9th and 10th rows based on line breaks
            apply_bullets = i >= 8  # 9th and 10th rows (index 8, 9)
            if '\n' in text_value:
                formatted_text = format_text_with_bullets(text_value, apply_bullets)
                lines = formatted_text.split('\n')
                
                if lines:
                    if apply_bullets:
                        # Add each line with first line without bullet, subsequent lines with bullets
                        for line_idx, line in enumerate(lines):
                            if line_idx > 0:
                                p = cell.add_paragraph(line)
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                p.paragraph_format.left_indent = Inches(0.25)
                                p.paragraph_format.space_before = Pt(2)
                                p.paragraph_format.space_after = Pt(2)
                                p.paragraph_format.line_spacing = 1.0
                            else:
                                p = cell.add_paragraph(line)
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                                p.paragraph_format.line_spacing = 1.0
                    else:
                        p = cell.add_paragraph(lines[0])
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        p.paragraph_format.line_spacing = 1.0
                        for line in lines[1:]:
                            p = cell.add_paragraph(line)
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.line_spacing = 1.0
            else:
                if apply_bullets and text_value:
                    formatted_text = format_text_with_bullets(text_value, apply_bullets)
                    p = cell.add_paragraph(formatted_text)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p.paragraph_format.line_spacing = 1.0
                else:
                    p = cell.add_paragraph(text_value)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    p.paragraph_format.line_spacing = 1.0
                value_run = p.runs[-1]
                value_run.font.size = Pt(11)
        
        row_index += 1
    
    # Add the last row with "Proof of Concept" and subsequent columns
    last_cell = table.cell(row_index, 0)
    last_paragraph = last_cell.paragraphs[0]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justify alignment
    last_paragraph.paragraph_format.line_spacing = 1.0
    set_cell_margins(last_cell, margin_value=100)
    
    # Add "Proof of Concept" header
    poc_header_run = last_paragraph.add_run("Proof of Concept:")
    poc_header_run.bold = True
    poc_header_run.font.size = Pt(12)
    if severity_hex:
        r = int(severity_hex[0:2], 16)
        g = int(severity_hex[2:4], 16)
        b = int(severity_hex[4:6], 16)
        poc_header_run.font.color.rgb = RGBColor(r, g, b)
    
    # Process all columns from "Proof of Concept" onwards
    step_counter = 1  # To track step numbers if a column doesn't start with "Step"
    for col_idx, col_name in enumerate(additional_columns):
        col_value = str(row.get(col_name, '')).strip()
        print(f"Row {idx + 1} - Column {col_name}: '{col_value}'")
        
        # Check if the column value contains a step or an image
        is_image = any(col_value.lower().endswith(ext) for ext in ['.png', '.jpg', '.jpeg'])
        is_step = re.search(r'Step\s*\d+:|\bstep\b', col_value, flags=re.IGNORECASE)
        
        if col_idx == 0:  # "Proof of Concept" column
            if col_value and col_value.lower() != 'nan':
                # Parse POC text into steps
                steps = re.split(r'(Step\s*\d+:)', col_value, flags=re.IGNORECASE)
                step_dict = {}
                current_step = None
                for part in steps:
                    part = part.strip()
                    if re.match(r'Step\s*\d+:', part, flags=re.IGNORECASE):
                        current_step = part
                        step_dict[current_step] = ""
                    elif current_step and part:
                        step_dict[current_step] += part
                
                print(f"Row {idx + 1} - Parsed Steps from POC: {step_dict}")
                
                if not step_dict:
                    # If no steps found, treat the entire POC text as a single step
                    step_name = f"Step{step_counter}:"
                    last_cell.add_paragraph()
                    step_para = last_cell.paragraphs[-1]
                    step_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    step_run = step_para.add_run(step_name)
                    step_run.bold = True
                    step_run.font.size = Pt(11)
                    
                    last_cell.add_paragraph()
                    content_para = last_cell.paragraphs[-1]
                    content_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    content_run = content_para.add_run(col_value)
                    content_run.font.size = Pt(11)
                    content_para.paragraph_format.space_before = Pt(2)
                    content_para.paragraph_format.space_after = Pt(2)
                    step_counter += 1
                else:
                    for step_name, step_content in step_dict.items():
                        # Add step name
                        last_cell.add_paragraph()
                        step_para = last_cell.paragraphs[-1]
                        step_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        step_run = step_para.add_run(step_name)
                        step_run.bold = True
                        step_run.font.size = Pt(11)
                        
                        # Add step content
                        formatted_content = format_text_with_bullets(step_content, apply_bullets=False)
                        last_cell.add_paragraph()
                        content_para = last_cell.paragraphs[-1]
                        content_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        content_run = content_para.add_run(formatted_content)
                        content_run.font.size = Pt(11)
                        content_para.paragraph_format.space_before = Pt(2)
                        content_para.paragraph_format.space_after = Pt(2)
                        step_counter += 1
        elif is_step:
            # Parse the column value as a step
            steps = re.split(r'(Step\s*\d+:)', col_value, flags=re.IGNORECASE)
            step_dict = {}
            current_step = None
            for part in steps:
                part = part.strip()
                if re.match(r'Step\s*\d+:', part, flags=re.IGNORECASE):
                    current_step = part
                    step_dict[current_step] = ""
                elif current_step and part:
                    step_dict[current_step] += part
            
            print(f"Row {idx + 1} - Parsed Steps from {col_name}: {step_dict}")
            
            if not step_dict:
                # If no "StepX:" found, treat the entire value as a single step
                step_name = f"Step{step_counter}:"
                last_cell.add_paragraph()
                step_para = last_cell.paragraphs[-1]
                step_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                step_run = step_para.add_run(step_name)
                step_run.bold = True
                step_run.font.size = Pt(11)
                
                last_cell.add_paragraph()
                content_para = last_cell.paragraphs[-1]
                content_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                content_run = content_para.add_run(col_value)
                content_run.font.size = Pt(11)
                content_para.paragraph_format.space_before = Pt(2)
                content_para.paragraph_format.space_after = Pt(2)
                step_counter += 1
            else:
                for step_name, step_content in step_dict.items():
                    # Add step name
                    last_cell.add_paragraph()
                    step_para = last_cell.paragraphs[-1]
                    step_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    step_run = step_para.add_run(step_name)
                    step_run.bold = True
                    step_run.font.size = Pt(11)
                    
                    # Add step content
                    formatted_content = format_text_with_bullets(step_content, apply_bullets=False)
                    last_cell.add_paragraph()
                    content_para = last_cell.paragraphs[-1]
                    content_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    content_run = content_para.add_run(formatted_content)
                    content_run.font.size = Pt(11)
                    content_para.paragraph_format.space_before = Pt(2)
                    content_para.paragraph_format.space_after = Pt(2)
                    step_counter += 1
        elif is_image:
            # Handle image column: display the actual image
            paths = [p.strip() for p in col_value.split(",") if p.strip()]
            for path in paths:
                if os.path.exists(path) and path.lower().endswith(('.png', '.jpg', '.jpeg')):
                    last_cell.add_paragraph()
                    img_para = last_cell.paragraphs[-1]
                    img_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    run = img_para.add_run()
                    run.add_picture(path, width=Inches(5.0))
                else:
                    print(f"Row {idx + 1} - Skipped image: {path}")
        else:
            # Treat as a step if it's not an image
            step_name = f"Step{step_counter}:"
            last_cell.add_paragraph()
            step_para = last_cell.paragraphs[-1]
            step_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            step_run = step_para.add_run(step_name)
            step_run.bold = True
            step_run.font.size = Pt(11)
            
            last_cell.add_paragraph()
            content_para = last_cell.paragraphs[-1]
            content_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            content_run = content_para.add_run(col_value)
            content_run.font.size = Pt(11)
            content_para.paragraph_format.space_before = Pt(2)
            content_para.paragraph_format.space_after = Pt(2)
            step_counter += 1
    
    # Apply table borders (with custom handling for first and second rows)
    set_table_borders(table)
    
    # Add page break
    doc.add_page_break()

# Step 6: Reattach trailing content
for elem in following_elements:
    parent.append(elem)

# Step 7: Save document
doc.save(output_file)
print(f"Document saved as {output_file}")

# Step 8: Run optional next script
next_script = "table_update.py"
try:
    result = subprocess.run([sys.executable, next_script], check=True, capture_output=True, text=True)
    print(f"Successfully ran {next_script}")
except subprocess.CalledProcessError as e:
    print(f"Error running {next_script}: {e}")
    print(f"Standard Output: {e.stdout}")
    print(f"Standard Error: {e.stderr}")
except FileNotFoundError:
    print(f"Error: {next_script} not found in the current directory.")